#!/usr/bin/env python3
"""
ScaleFlow Branded Document Generator v2

Generates visually branded .docx files using full-width color blocks,
dark title headers, bold section strips, and proper brand presence.

Usage:
    python3 shared/generate_branded_docx.py --input data.json --output output.docx

JSON input structure: see README or Brief Analyzer SKILL.md for full schema.
"""

import json
import sys
import os
import argparse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# COLOR UTILITIES
# ============================================================

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def hex_to_str(hex_color):
    return hex_color.lstrip("#").upper()

def lighten_hex(hex_color, factor=0.92):
    hex_color = hex_color.lstrip("#")
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    r = int(r + (255 - r) * factor)
    g = int(g + (255 - g) * factor)
    b = int(b + (255 - b) * factor)
    return f"{r:02X}{g:02X}{b:02X}"

def darken_hex(hex_color, factor=0.3):
    hex_color = hex_color.lstrip("#")
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    r = int(r * (1 - factor))
    g = int(g * (1 - factor))
    b = int(b * (1 - factor))
    return f"{r:02X}{g:02X}{b:02X}"


# ============================================================
# FONT MAPPING — always use system-safe fallbacks
# ============================================================

FONT_FALLBACKS = {
    "Bebas Neue": "Arial Black",
    "Inter": "Calibri",
    "Montserrat": "Arial",
    "Playfair Display": "Georgia",
    "Roboto": "Calibri",
    "DM Sans": "Calibri",
    "Poppins": "Calibri",
    "Helvetica Neue": "Helvetica",
    "Garamond": "Garamond",
}

STYLE_FONT_MAP = {
    "modern": ("Arial", "Calibri"),
    "classic": ("Georgia", "Georgia"),
    "bold": ("Arial Black", "Arial"),
    "minimal": ("Helvetica", "Calibri"),
    "editorial": ("Garamond", "Garamond"),
}


# ============================================================
# XML HELPERS
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, color="D0D0D0", width="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for name in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), width)
        b.set(qn("w:color"), color)
        b.set(qn("w:space"), "0")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for name in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

def set_row_height(row, height_pt):
    trPr = row._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(height_pt * 20)))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)

def set_paragraph_spacing(p, before=0, after=0, line=None):
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

def make_full_width_block(doc, bg_color, height_pt=None):
    """Create a full-width single-cell table that acts as a color block.
    Returns the cell so you can add content to it."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Make table full width
    tblPr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)
    set_no_borders(cell)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

    if height_pt:
        set_row_height(table.rows[0], height_pt)

    # Remove default empty paragraph
    cell.text = ""

    return cell, table


# ============================================================
# DOCUMENT BUILDER v2
# ============================================================

class BrandedDocBuilder:
    def __init__(self, brand_config):
        self.brand = brand_config
        self.doc = Document()

        self.primary = hex_to_rgb(brand_config["primary_color"])
        self.secondary = hex_to_rgb(brand_config["secondary_color"])
        self.accent = hex_to_rgb(brand_config["accent_color"])
        self.primary_hex = hex_to_str(brand_config["primary_color"])
        self.secondary_hex = hex_to_str(brand_config["secondary_color"])
        self.accent_hex = hex_to_str(brand_config["accent_color"])
        self.light_primary = lighten_hex(brand_config["primary_color"], 0.92)
        self.light_accent = lighten_hex(brand_config["accent_color"], 0.93)
        self.dark_secondary = darken_hex(brand_config["secondary_color"], 0.2)

        # Resolve fonts — always use system-safe fallback
        h_font = brand_config.get("heading_font", "")
        b_font = brand_config.get("body_font", "")
        style = brand_config.get("typography_style", "modern")

        if h_font and h_font != "not specified":
            self.heading_font = FONT_FALLBACKS.get(h_font, h_font)
        else:
            self.heading_font = STYLE_FONT_MAP.get(style, ("Arial", "Calibri"))[0]

        if b_font and b_font != "not specified":
            self.body_font = FONT_FALLBACKS.get(b_font, b_font)
        else:
            self.body_font = STYLE_FONT_MAP.get(style, ("Arial", "Calibri"))[1]

        self.white = RGBColor(0xFF, 0xFF, 0xFF)
        self.body_color = RGBColor(0x2A, 0x2A, 0x2A)
        self.subtle_color = RGBColor(0x77, 0x77, 0x77)

        # Check if logo exists
        logo_path = brand_config.get("logo_path")
        self.logo_path = logo_path if logo_path and os.path.exists(logo_path) else None

        self._setup_page()

    def _setup_page(self):
        for section in self.doc.sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

    def _run(self, p, text, font=None, size=None, color=None, bold=False, italic=False):
        run = p.add_run(text)
        run.font.name = font or self.body_font
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
        return run

    # ----------------------------------------------------------
    # TITLE BLOCK — Dark full-width header with brand colors
    # ----------------------------------------------------------
    def add_title_block(self, title, subtitle, date, meta_pairs):
        # === HERO HEADER: Full-width dark block with title ===
        cell, _ = make_full_width_block(self.doc, self.secondary_hex)
        set_cell_margins(cell, top=350, bottom=100, left=250, right=250)

        # Accent stripe inside header
        accent_p = cell.paragraphs[0]
        set_paragraph_spacing(accent_p, after=180)
        run = accent_p.add_run("━" * 40)
        run.font.color.rgb = self.accent
        run.font.size = Pt(8)
        run.font.name = self.body_font

        # Title text
        title_p = cell.add_paragraph()
        set_paragraph_spacing(title_p, after=60)
        self._run(title_p, title, font=self.heading_font, size=28, color=self.primary, bold=True)

        # Subtitle
        sub_p = cell.add_paragraph()
        set_paragraph_spacing(sub_p, after=40)
        self._run(sub_p, subtitle, size=12, color=self.white)

        # Date
        date_p = cell.add_paragraph()
        set_paragraph_spacing(date_p, after=100)
        self._run(date_p, date, size=10, color=self.subtle_color)

        # Logo if available
        if self.logo_path:
            try:
                logo_p = cell.add_paragraph()
                logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = logo_p.add_run()
                run.add_picture(self.logo_path, width=Inches(1.5))
            except Exception:
                pass

        # === ACCENT BAR ===
        accent_cell, _ = make_full_width_block(self.doc, self.accent_hex)
        set_cell_margins(accent_cell, top=0, bottom=0, left=0, right=0)
        set_row_height(_.rows[0], 5)

        # === META INFO BLOCK ===
        if meta_pairs:
            self.doc.add_paragraph()  # small spacer
            meta_table = self.doc.add_table(rows=len(meta_pairs), cols=2)
            meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Full width
            tblPr = meta_table._tbl.tblPr
            tblW = OxmlElement("w:tblW")
            tblW.set(qn("w:w"), "5000")
            tblW.set(qn("w:type"), "pct")
            tblPr.append(tblW)

            for i, pair in enumerate(meta_pairs):
                # Label
                lc = meta_table.rows[i].cells[0]
                lc.text = ""
                p = lc.paragraphs[0]
                self._run(p, pair["label"].upper(), size=8, color=self.subtle_color, bold=True)
                set_cell_shading(lc, "F7F7F7")
                set_cell_borders(lc, "EEEEEE", "2")
                set_cell_margins(lc, top=40, bottom=40, left=120, right=60)

                # Value
                vc = meta_table.rows[i].cells[1]
                vc.text = ""
                p = vc.paragraphs[0]
                self._run(p, pair["value"], size=9, color=self.body_color)
                set_cell_shading(vc, "FAFAFA")
                set_cell_borders(vc, "EEEEEE", "2")
                set_cell_margins(vc, top=40, bottom=40, left=120, right=60)

            self.doc.add_paragraph()  # spacer after meta

    # ----------------------------------------------------------
    # SECTION HEADING — Full-width primary color strip
    # ----------------------------------------------------------
    def add_section_heading(self, text):
        cell, _ = make_full_width_block(self.doc, self.primary_hex)
        set_cell_margins(cell, top=100, bottom=100, left=200, right=200)
        set_row_height(_.rows[0], 28)

        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        self._run(p, text.upper(), font=self.heading_font, size=12, color=self.white, bold=True)

    # ----------------------------------------------------------
    # CONTENT BLOCKS
    # ----------------------------------------------------------
    def add_paragraph(self, text, bold=False, italic=False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        self._run(p, text, size=10, color=self.body_color, bold=bold, italic=italic)
        return p

    def add_bullets(self, items):
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.8)

            # Accent-colored bullet marker
            self._run(p, "●  ", size=8, color=self.accent, bold=True)
            self._run(p, item, size=10, color=self.body_color)

    def add_labeled_pairs(self, pairs):
        """Labeled pairs as a clean two-column mini-table."""
        table = self.doc.add_table(rows=len(pairs), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        tblPr = table._tbl.tblPr
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        tblPr.append(tblW)

        for i, pair in enumerate(pairs):
            lc = table.rows[i].cells[0]
            lc.text = ""
            p = lc.paragraphs[0]
            self._run(p, pair["label"], size=9, color=self.body_color, bold=True)
            set_no_borders(lc)
            set_cell_margins(lc, top=30, bottom=30, left=100, right=40)

            # Left accent border on label cell
            tc = lc._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            left_b = OxmlElement("w:left")
            left_b.set(qn("w:val"), "single")
            left_b.set(qn("w:sz"), "16")
            left_b.set(qn("w:color"), self.accent_hex)
            left_b.set(qn("w:space"), "0")
            tcBorders.append(left_b)
            tcPr.append(tcBorders)

            vc = table.rows[i].cells[1]
            vc.text = ""
            p = vc.paragraphs[0]
            self._run(p, pair["value"], size=9, color=self.subtle_color)
            set_no_borders(vc)
            set_cell_margins(vc, top=30, bottom=30, left=40, right=100)

        self.doc.add_paragraph()  # spacer

    def add_table(self, headers, rows):
        """Branded data table with strong header row."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        tblPr = table._tbl.tblPr
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        tblPr.append(tblW)

        # Header row — secondary (dark) background
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            self._run(p, header, size=9, color=self.white, bold=True)
            set_cell_shading(cell, self.secondary_hex)
            set_cell_borders(cell, self.secondary_hex, "2")
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)

        # Data rows
        for r, row_data in enumerate(rows):
            for c, value in enumerate(row_data):
                cell = table.rows[r + 1].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                self._run(p, str(value), size=9, color=self.body_color)
                # Alternate rows with light primary tint
                bg = self.light_primary if r % 2 == 0 else "FFFFFF"
                set_cell_shading(cell, bg)
                set_cell_borders(cell, "E0E0E0", "2")
                set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

        self.doc.add_paragraph()  # spacer

    def add_callout(self, text):
        """Callout box with strong accent left border and tinted background."""
        cell, tbl = make_full_width_block(self.doc, self.light_accent)
        set_cell_margins(cell, top=100, bottom=100, left=250, right=200)

        # Strong left border in accent color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        left_b = OxmlElement("w:left")
        left_b.set(qn("w:val"), "single")
        left_b.set(qn("w:sz"), "36")
        left_b.set(qn("w:color"), self.accent_hex)
        left_b.set(qn("w:space"), "0")
        tcBorders.append(left_b)
        for name in ["top", "right", "bottom"]:
            b = OxmlElement(f"w:{name}")
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            tcBorders.append(b)
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        self._run(p, text, size=9, color=self.body_color, italic=True)

    # ----------------------------------------------------------
    # FOOTER
    # ----------------------------------------------------------
    def add_footer(self, text):
        # Accent bar
        accent_cell, tbl = make_full_width_block(self.doc, self.accent_hex)
        set_cell_margins(accent_cell, top=0, bottom=0, left=0, right=0)
        set_row_height(tbl.rows[0], 3)

        # Secondary bar
        sec_cell, tbl2 = make_full_width_block(self.doc, self.secondary_hex)
        set_cell_margins(sec_cell, top=80, bottom=80, left=200, right=200)

        p = sec_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=0)
        self._run(p, text, size=8, color=self.subtle_color)

    # ----------------------------------------------------------
    # PAGE HEADER/FOOTER (repeating on every page)
    # ----------------------------------------------------------
    def add_page_header_footer(self, client_name, campaign_name):
        """Add repeating header and footer to every page."""
        for section in self.doc.sections:
            # Header — thin accent line + client name
            header = section.header
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = hp.add_run(f"{client_name}  |  {campaign_name}")
            run.font.name = self.body_font
            run.font.size = Pt(7)
            run.font.color.rgb = self.subtle_color

            # Add bottom border to header paragraph
            pPr = hp._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:color"), self.accent_hex)
            bottom.set(qn("w:space"), "4")
            pBdr.append(bottom)
            pPr.append(pBdr)

            # Footer — page context
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run("Prepared by ScaleFlow")
            run.font.name = self.body_font
            run.font.size = Pt(7)
            run.font.color.rgb = self.subtle_color

    # ----------------------------------------------------------
    # BUILD FROM JSON DATA
    # ----------------------------------------------------------
    def build_from_data(self, doc_data):
        # Page headers/footers
        client = self.brand.get("client_name", "")
        campaign = doc_data.get("title", "")
        self.add_page_header_footer(client, campaign)

        # Title block
        self.add_title_block(
            title=doc_data["title"],
            subtitle=doc_data.get("subtitle", ""),
            date=doc_data.get("date", ""),
            meta_pairs=doc_data.get("meta", []),
        )

        # Sections
        for section in doc_data.get("sections", []):
            self.add_section_heading(section["heading"])

            for content in section.get("content", []):
                ctype = content["type"]
                if ctype == "paragraph":
                    self.add_paragraph(
                        content["text"],
                        bold=content.get("bold", False),
                        italic=content.get("italic", False),
                    )
                elif ctype == "bullets":
                    self.add_bullets(content["items"])
                elif ctype == "labeled":
                    self.add_labeled_pairs(content["pairs"])
                elif ctype == "table":
                    self.add_table(content["headers"], content["rows"])
                elif ctype == "callout":
                    self.add_callout(content["text"])

        # Document footer
        footer_text = doc_data.get("footer", "")
        if footer_text:
            self.add_footer(footer_text)

    def save(self, output_path):
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        self.doc.save(output_path)
        print(f"Branded document saved: {output_path}")
        return output_path


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="Generate a branded .docx from JSON data.")
    parser.add_argument("--input", required=True, help="Path to JSON input file")
    parser.add_argument("--output", required=True, help="Path for output .docx file")
    args = parser.parse_args()

    with open(args.input, "r") as f:
        data = json.load(f)

    builder = BrandedDocBuilder(data["brand"])
    builder.build_from_data(data["document"])
    builder.save(args.output)


if __name__ == "__main__":
    main()
